#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate formatted Word document from thesis markdown.
Wuhan University of Technology undergraduate thesis formatting.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ===== Configuration =====
MARKDOWN_PATH = r'd:\SwapU\docs\thesis\初稿_润色.md'
OUTPUT_PATH = r'd:\SwapU\docs\thesis\SwapU毕业论文.docx'

# Font names
FONT_HEI = '黑体'
FONT_SONG = '宋体'
FONT_TNR = 'Times New Roman'

# Font sizes (in Pt)
SIZE_ER = Pt(18)       # 小二 18pt - 章标题, 摘要标题
SIZE_SAN = Pt(15)      # 小三 15pt - 一级节标题
SIZE_SI = Pt(14)       # 四号 14pt - 二级节标题, 关键词标题
SIZE_XIAOSI = Pt(12)   # 小四 12pt - 正文, 三级节标题
SIZE_WU = Pt(10.5)     # 五号 10.5pt - 参考文献正文
SIZE_XIAOWU = Pt(9)    # 小五 9pt - 页眉页脚

# Page margins
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)

# Line spacing (fixed 20pt)
LINE_SPACING = Pt(20)

# ===== Helper Functions =====

def set_cell_font(run, font_cn, font_en, size, bold=False):
    """Set both Chinese and English font on a run."""
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)

def set_paragraph_spacing(paragraph, line_spacing=LINE_SPACING, space_before=0, space_after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after

def add_title_paragraph(doc, text, font_cn, font_en, size, alignment, space_before, space_after, bold=True):
    """Add a title paragraph with formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    set_paragraph_spacing(p, LINE_SPACING, space_before, space_after)
    run = p.add_run(text)
    set_cell_font(run, font_cn, font_en, size, bold)
    return p

def add_body_paragraph(doc, text, first_line_indent=True):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, LINE_SPACING, 0, 0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 Chinese chars
    run = p.add_run(text)
    set_cell_font(run, FONT_SONG, FONT_TNR, SIZE_XIAOSI, False)
    return p

def add_chapter_title(doc, text):
    """Add chapter-level title (第X章 ...)."""
    return add_title_paragraph(doc, text, FONT_HEI, FONT_TNR, SIZE_ER,
                               WD_ALIGN_PARAGRAPH.CENTER, Pt(10), Pt(10), True)

def add_section1_title(doc, text):
    """Add level-1 section title (X.X ...)."""
    return add_title_paragraph(doc, text, FONT_HEI, FONT_TNR, SIZE_SAN,
                               WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(6), True)

def add_section2_title(doc, text):
    """Add level-2 section title."""
    return add_title_paragraph(doc, text, FONT_HEI, FONT_TNR, SIZE_SI,
                               WD_ALIGN_PARAGRAPH.LEFT, Pt(4), Pt(4), True)

def add_section3_title(doc, text):
    """Add level-3 section title (if any)."""
    return add_title_paragraph(doc, text, FONT_HEI, FONT_TNR, SIZE_XIAOSI,
                               WD_ALIGN_PARAGRAPH.LEFT, Pt(2), Pt(2), True)

def add_figure_caption(doc, text):
    """Add figure/table caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, Pt(16), 0, 0)
    run = p.add_run(text)
    set_cell_font(run, FONT_SONG, FONT_TNR, Pt(10.5), False)
    return p

def add_mermaid_placeholder(doc, fig_num, fig_title):
    """Add placeholder for Mermaid diagram."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, Pt(16), 4, 4)
    run = p.add_run(f'[图 {fig_num} 图表 - 请手动插入]')
    set_cell_font(run, FONT_SONG, FONT_TNR, Pt(9), False)
    # Caption
    cap = add_figure_caption(doc, f'图 {fig_num} {fig_title}')
    return p

def add_reference(doc, text):
    """Add reference entry."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, Pt(16), 0, 0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_cell_font(run, FONT_SONG, FONT_TNR, SIZE_WU, False)
    return p

def setup_page(doc):
    """Configure page settings for all sections."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

def add_header_footer(doc, thesis_title_short):
    """Add header and footer to sections starting from body."""
    for i, section in enumerate(doc.sections):
        if i == 0:
            # First section (abstract/TOC) - no header/footer needed
            # or use Roman numerals for page numbers
            continue

        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(hp, Pt(12), 0, 0)
        run = hp.add_run(f'武汉理工大学学士学位论文{thesis_title_short}')
        set_cell_font(run, FONT_SONG, FONT_TNR, SIZE_XIAOWU, False)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)
        set_cell_font(run, FONT_TNR, FONT_TNR, SIZE_XIAOWU, False)
        set_cell_font(run2, FONT_TNR, FONT_TNR, SIZE_XIAOWU, False)
        set_cell_font(run3, FONT_TNR, FONT_TNR, SIZE_XIAOWU, False)


# ===== Main =====

def main():
    # Read markdown
    with open(MARKDOWN_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # ---- Set default style ----
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_SONG
    font.size = SIZE_XIAOSI
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

    # ---- Parse and build document ----
    lines = content.split('\n')
    i = 0
    fig_counter = 0

    # Phase tracking
    in_abstract_cn = False
    in_abstract_en = False
    in_keywords_cn = False
    in_keywords_en = False
    in_toc = False
    in_mermaid = False
    in_body = False
    in_references = False
    current_chapter = 0

    # Track reference numbers for chapters 1-7
    chapter_refs = {
        1: ['[1] 徐雅靖. 以提升信任为导向的二手交易APP的体验设计研究[D]. 华东理工大学, 2022.',
            '[2] 王颖. 社区二手物品处置系统研究与设计[D]. 哈尔滨理工大学, 2022.',
            '[3] 庞媛媛. 闲置交易平台颠覆式创新路径研究[D]. 北京邮电大学, 2025.',
            '[4] 王祎. 基于RAG技术的红色档案智能问答系统构建与应用[D]. 上海大学, 2025.',
            '[5] 杨梓鹏. 基于RAG的多语言跨境电商智能客服应用研究[D]. 广东财经大学, 2025.']
    }

    body_started = False
    abstract_text_cn = []
    abstract_text_en = []

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Detect mermaid blocks
        if line.startswith('```mermaid'):
            in_mermaid = True
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            in_mermaid = False
            i += 1  # Skip closing ```
            # Skip the caption line too if it's a <p align> tag
            if i < len(lines) and lines[i].strip().startswith('<p'):
                # Extract figure caption
                cap_match = re.search(r'<b>(.*?)</b>', lines[i])
                if cap_match:
                    fig_counter += 1
                    add_mermaid_placeholder(doc, f'{current_chapter}-{fig_counter}', cap_match.group(1))
                i += 1
            continue

        # Skip HTML caption lines
        if line.startswith('<p'):
            i += 1
            continue

        # Detect Chinese abstract
        if line == '# **摘要**':
            in_abstract_cn = True
            add_chapter_title(doc, '摘要')
            i += 1
            continue

        # Detect English abstract
        if line == '# **Abstract**':
            in_abstract_cn = False
            in_abstract_en = True
            # Add page break before English abstract
            doc.add_page_break()
            add_chapter_title(doc, 'Abstract')
            i += 1
            continue

        # Detect keywords in Chinese abstract
        if in_abstract_cn and '关键词' in line:
            # Add keywords - aggressively clean markdown formatting
            kw_text = line.strip()
            kw_text = re.sub(r'\\\*+', '', kw_text)  # Remove \* escaped asterisks
            kw_text = re.sub(r'\*+', '', kw_text)     # Remove all remaining asterisks
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, LINE_SPACING, 0, 0)
            # "关键词" label in bold
            if '：' in kw_text:
                label, values = kw_text.split('：', 1)
            elif ':' in kw_text:
                label, values = kw_text.split(':', 1)
            else:
                label, values = kw_text, ''
            run_label = p.add_run(label + '：')
            set_cell_font(run_label, FONT_HEI, FONT_TNR, SIZE_SI, True)
            run_values = p.add_run(values)
            set_cell_font(run_values, FONT_SONG, FONT_TNR, SIZE_XIAOSI, False)
            i += 1
            continue

        # Detect keywords in English abstract
        if in_abstract_en and 'Keywords' in line:
            kw_text = line.strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, LINE_SPACING, 0, 0)
            if ':' in kw_text:
                label, values = kw_text.split(':', 1)
                run_label = p.add_run(label + ': ')
                set_cell_font(run_label, FONT_TNR, FONT_TNR, SIZE_SI, True)
                run_values = p.add_run(values.strip())
                set_cell_font(run_values, FONT_TNR, FONT_TNR, SIZE_XIAOSI, False)
            i += 1
            continue

        # Detect chapter titles
        chapter_match = re.match(r'^# \*\*第(\d+)章\s+(.+?)\*\*$', line)
        if chapter_match:
            in_abstract_cn = False
            in_abstract_en = False
            current_chapter = int(chapter_match.group(1))
            chapter_title = chapter_match.group(2).strip()
            full_title = f'第{current_chapter}章 {chapter_title}'

            if not body_started:
                body_started = True
                # Add section break for header/footer
                doc.add_section()
                setup_page(doc)
                add_header_footer(doc, '')

            doc.add_page_break()
            add_chapter_title(doc, full_title)
            i += 1
            continue

        # Detect section titles (## X.X ...)
        section_match = re.match(r'^##\s+(\d+\.\d+)\s*(.*)', line)
        if section_match:
            sec_num = section_match.group(1)
            sec_title = section_match.group(2).strip()
            full_sec = f'{sec_num} {sec_title}'
            # Determine section level
            if sec_num.count('.') == 1:
                add_section1_title(doc, full_sec)
            elif sec_num.count('.') == 2:
                add_section2_title(doc, full_sec)
            elif sec_num.count('.') == 3:
                add_section3_title(doc, full_sec)
            i += 1
            continue

        # Skip TOC entries (lines with only #_Toc references)
        if 'Toc' in line and ('_Toc' in line or '#_Toc' in line):
            i += 1
            continue

        # Skip pure markdown formatting lines
        if line in ['---', '***', '___'] or line.startswith('***'):
            i += 1
            continue

        # Handle markdown tables (must be before body text cleaning)
        if line.startswith('|') and line.endswith('|'):
            if '---' in line:
                i += 1
                continue  # Skip separator row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if 'table_active' not in dir(doc):
                doc.table_active = True
                doc.current_table = doc.add_table(rows=1, cols=len(cells))
                doc.current_table.style = 'Table Grid'
                for j, cell_text in enumerate(cells):
                    cell = doc.current_table.rows[0].cells[j]
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(cell_text)
                    set_cell_font(run, FONT_HEI, FONT_TNR, Pt(9), True)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row = doc.current_table.add_row()
                for j, cell_text in enumerate(cells):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = ''
                        run = cell.paragraphs[0].add_run(cell_text)
                        set_cell_font(run, FONT_SONG, FONT_TNR, Pt(9), False)
            i += 1
            continue

        # Detect end of table
        if 'table_active' in dir(doc) and not line.startswith('|'):
            delattr(doc, 'table_active')
            delattr(doc, 'current_table')
            # fall through to process this line normally
        if line.startswith('> [') and ']' in line:
            ref_text = line.lstrip('> ').strip()
            add_reference(doc, ref_text)
            i += 1
            continue

        # Handle bold reference labels like **[1]**
        if line.startswith('**[1]**') or line.startswith('> **[1]**'):
            ref_text = re.sub(r'\*\*|\> ', '', line).strip()
            add_reference(doc, ref_text)
            i += 1
            continue

        # Regular body text - clean markdown formatting
        clean_line = line
        # Remove escaped asterisks and bold-italic patterns (***\*text\**** -> text)
        clean_line = re.sub(r'\*{1,3}\\\*', '', clean_line)
        clean_line = re.sub(r'\\\*\*{1,3}', '', clean_line)
        # Remove bold+italic markers
        clean_line = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', clean_line)
        # Remove bold markers (**text** -> text)
        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)
        # Remove italic markers (*text* -> text, but not when part of **)
        clean_line = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', clean_line)
        # Remove inline code markers
        clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)
        # Remove HTML tags
        clean_line = re.sub(r'<[^>]+>', '', clean_line)
        # Remove leftover standalone backslash-asterisks
        clean_line = re.sub(r'\\\*+', '', clean_line)
        clean_line = re.sub(r'\*+\\', '', clean_line)
        # Remove residual *** patterns (bold-italic markers mixed with backslashes)
        clean_line = re.sub(r'\*{2,}', '', clean_line)

        if clean_line.strip():
            # Determine context
            if in_abstract_cn:
                add_body_paragraph(doc, clean_line, first_line_indent=False)
            elif in_abstract_en:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_spacing(p, LINE_SPACING, 0, 0)
                run = p.add_run(clean_line)
                set_cell_font(run, FONT_TNR, FONT_TNR, SIZE_XIAOSI, False)
            elif not body_started:
                # Pre-body content (before chapter 1)
                pass
            else:
                # Check if it's a reference in body text
                if clean_line.startswith('**参考文献') or clean_line.startswith('参考文献'):
                    add_chapter_title(doc, '参考文献')
                else:
                    add_body_paragraph(doc, clean_line)

        i += 1

    # ---- Save ----
    setup_page(doc)
    doc.save(OUTPUT_PATH)
    print(f'Document saved to: {OUTPUT_PATH}')
    print(f'Total figures: {fig_counter}')

if __name__ == '__main__':
    main()
