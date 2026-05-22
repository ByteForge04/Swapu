import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v11.docx')

print("=" * 60)
print("SECTIONS - HEADER/FOOTER DETAILS")
print("=" * 60)
for i, sec in enumerate(doc.sections):
    print(f"\nSection {i}:")
    header = sec.header
    if header:
        print(f"  Header linked_to_previous: {header.is_linked_to_previous}")
        for j, hp in enumerate(header.paragraphs):
            if hp.text.strip() or hp.runs:
                print(f"  Header P[{j}]: '{hp.text}'")
                for run in hp.runs:
                    font = run.font
                    rPr = run._element.find(qn('w:rPr'))
                    ea_font = ""
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ea_font = rFonts.get(qn('w:eastAsia'), '')
                    print(f"    Run: '{run.text}' font={font.name}/{ea_font} size={font.size} bold={font.bold}")
    
    footer = sec.footer
    if footer:
        print(f"  Footer linked_to_previous: {footer.is_linked_to_previous}")
        for j, fp in enumerate(footer.paragraphs):
            if fp.text.strip() or fp.runs:
                print(f"  Footer P[{j}]: '{fp.text}'")
                for run in fp.runs:
                    font = run.font
                    rPr = run._element.find(qn('w:rPr'))
                    ea_font = ""
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ea_font = rFonts.get(qn('w:eastAsia'), '')
                    print(f"    Run: '{run.text}' font={font.name}/{ea_font} size={font.size} bold={font.bold}")
            for child in fp._element:
                if child.tag == qn('w:r'):
                    for sub in child:
                        if sub.tag in (qn('w:fldChar'), qn('w:instrText')):
                            val = sub.text if sub.text else sub.get(qn('w:fldCharType'), '')
                            print(f"    Field: {sub.tag.split('}')[1]} = {val}")

    # Check section type (section break)
    sectPr = sec._sectPr
    if sectPr is not None:
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print(f"  pgNumType: fmt={fmt} start={start}")
        else:
            print(f"  pgNumType: None")

print("\n" + "=" * 60)
print("LAST 10 PARAGRAPHS")
print("=" * 60)
total = len(doc.paragraphs)
for i in range(max(0, total - 10), total):
    p = doc.paragraphs[i]
    text = p.text.strip()
    style_name = p.style.name if p.style else "None"
    font_info = ""
    if p.runs:
        r = p.runs[0]
        font = r.font
        rPr = r._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        font_info = f"font={font.name}/{ea_font} size={font.size} bold={font.bold}"
    display_text = text[:80] if text else "(empty)"
    print(f"P[{i}] style='{style_name}' {font_info} | {display_text}")

print("\n" + "=" * 60)
print("LINE SPACING CHECK (sample paragraphs)")
print("=" * 60)
for i in [30, 43, 47, 52, 69, 167, 203, 230, 241]:
    if i < total:
        p = doc.paragraphs[i]
        pf = p.paragraph_format
        ls = pf.line_spacing
        lsr = pf.line_spacing_rule
        sb = pf.space_before
        sa = pf.space_after
        print(f"P[{i}] line_spacing={ls} rule={lsr} space_before={sb} space_after={sa} | {p.text[:40]}")
