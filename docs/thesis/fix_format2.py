import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

def set_run_font(run, cn_font, en_font, size_pt, bold=None):
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

print("Fix 1: 关键词行 - 分离标签和内容...")
p33 = doc.paragraphs[33]
full_text = p33.text
for run in p33.runs:
    p33._element.remove(run._element)

label_run = p33.add_run('关键词：')
set_run_font(label_run, '黑体', 'Times New Roman', 14, bold=None)
content_run = p33.add_run('校园二手交易；检索增强生成；大语言模型；Spring Boot；Vue 3')
set_run_font(content_run, '宋体', 'Times New Roman', 12, bold=None)

print("Fix 2: Key Words行 - 分离标签和内容...")
p39 = doc.paragraphs[39]
full_text_en = p39.text
for run in p39.runs:
    p39._element.remove(run._element)

label_run_en = p39.add_run('Key Words: ')
set_run_font(label_run_en, 'Times New Roman', 'Times New Roman', 14, bold=True)
content_run_en = p39.add_run('Campus Second-hand Trading; Retrieval-Augmented Generation; Large Language Model; Spring Boot; Vue 3')
set_run_font(content_run_en, 'Times New Roman', 'Times New Roman', 12, bold=None)

print("Fix 3: 页眉字号改为宋体五号(10.5pt)...")
for sec in doc.sections:
    header = sec.header
    for hp in header.paragraphs:
        for run in hp.runs:
            set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=None)

print("Fix 4: 页码字号改为TNR五号(10.5pt)...")
for sec in doc.sections:
    footer = sec.footer
    for fp in footer.paragraphs:
        for run in fp.runs:
            if run.text and run.text.strip().isdigit():
                set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=None)

print("Fix 5: 摘要标题加空格 '摘    要'...")
p29 = doc.paragraphs[29]
for run in p29.runs:
    if '摘要' in run.text:
        run.text = '摘    要'

print("Fix 6: 致谢标题加空格 '致    谢'...")
p262 = doc.paragraphs[262]
for run in p262.runs:
    if '致谢' in run.text:
        run.text = '致    谢'

output_path = r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
