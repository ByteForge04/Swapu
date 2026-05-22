import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu
import json

doc = Document(r'd:\SwapU\要求和模板\附件1、武汉理工大学本科生毕业设计（论文）（模版）.docx')

print("=" * 60)
print("PAGE SETUP")
print("=" * 60)
for i, sec in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  page_width: {sec.page_width.cm:.2f} cm")
    print(f"  page_height: {sec.page_height.cm:.2f} cm")
    print(f"  top_margin: {sec.top_margin.cm:.2f} cm")
    print(f"  bottom_margin: {sec.bottom_margin.cm:.2f} cm")
    print(f"  left_margin: {sec.left_margin.cm:.2f} cm")
    print(f"  right_margin: {sec.right_margin.cm:.2f} cm")
    print(f"  header_distance: {sec.header_distance.cm:.2f} cm")
    print(f"  footer_distance: {sec.footer_distance.cm:.2f} cm")
    print(f"  gutter: {sec.gutter.cm:.2f} cm")
    print(f"  orientation: {sec.orientation}")
    print(f"  different_first_page_header_footer: {sec.different_first_page_header_footer}")
    
    # Check header
    header = sec.header
    if header and header.paragraphs:
        for hp in header.paragraphs:
            if hp.text.strip():
                print(f"  Header text: '{hp.text}'")
                for run in hp.runs:
                    font = run.font
                    print(f"    Run: '{run.text}' font_name={font.name} font_size={font.size} bold={font.bold}")
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            print(f"    rFonts: eastAsia={rFonts.get(qn('w:eastAsia'))} ascii={rFonts.get(qn('w:ascii'))} hAnsi={rFonts.get(qn('w:hAnsi'))}")
                pf = hp.paragraph_format
                print(f"    alignment={pf.alignment}")
    
    # Check footer / page number
    footer = sec.footer
    if footer and footer.paragraphs:
        for fp in footer.paragraphs:
            if fp.text.strip():
                print(f"  Footer text: '{fp.text}'")
            # Check for page number fields
            for run in fp.runs:
                fldChar = run._element.findall(qn('w:fldChar'))
                instrText = run._element.findall(qn('w:instrText'))
                if fldChar or instrText:
                    print(f"  Footer has page number field")
            for child in fp._element:
                if child.tag == qn('w:r'):
                    for sub in child:
                        if sub.tag == qn('w:fldChar') or sub.tag == qn('w:instrText'):
                            print(f"  Footer field: {sub.tag} = {sub.text if sub.text else sub.get(qn('w:fldCharType'), '')}")

print("\n" + "=" * 60)
print("STYLES")
print("=" * 60)
for style in doc.styles:
    if style.type is not None and hasattr(style, 'font') and style.font.size:
        font = style.font
        if font.size and font.size >= Pt(9):
            print(f"Style: '{style.name}' type={style.type} font_name={font.name} size={font.size} bold={font.bold}")

print("\n" + "=" * 60)
print("PARAGRAPHS (first 80)")
print("=" * 60)
for i, p in enumerate(doc.paragraphs):
    if i >= 80:
        break
    text = p.text.strip()
    if not text and not p.runs:
        continue
    style_name = p.style.name if p.style else "None"
    pf = p.paragraph_format
    alignment = pf.alignment
    
    # Get font info from first run
    font_info = ""
    if p.runs:
        r = p.runs[0]
        font = r.font
        rPr = r._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        font_info = f"font={font.name}/{ea_font} size={font.size} bold={font.bold}"
    
    # Check outline level
    outline_lvl = ""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl = f" outlineLvl={outlineLvl.get(qn('w:val'))}"
    
    display_text = text[:60] if text else "(empty)"
    print(f"P[{i}] style='{style_name}' align={alignment} {font_info}{outline_lvl} | {display_text}")

print("\n" + "=" * 60)
print("TABLES")
print("=" * 60)
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    if table.rows:
        first_row_text = [cell.text[:20] for cell in table.rows[0].cells]
        print(f"  First row: {first_row_text}")
