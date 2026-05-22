import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v14.docx')

keywords = ['并发', '超卖', '分布式锁', 'Redisson', '限流', 'rateLimiter', 'RateLimiter', 'RedissonClient', 'getLock', 'tryLock', '高并发', '防超卖', '并发购买', '并发场景', '并发下', '500 并发', '500并发', 'JMeter', 'jmeter', '压测', '施压']

print("=" * 80)
print("VERIFICATION: SEARCHING FOR REMAINING CONCURRENCY/ANTI-OVERSELLING CONTENT")
print("=" * 80)

found_any = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    found = []
    for kw in keywords:
        if kw in text:
            found.append(kw)
    if found:
        found_any = True
        print(f"\nP[{i}] keywords: {found}")
        print(f"  Text: {text[:200]}")

print("\n" + "=" * 80)
print("SEARCHING IN TABLES")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if not text:
                continue
            found = []
            for kw in keywords:
                if kw in text:
                    found.append(kw)
            if found:
                found_any = True
                print(f"\nTable[{ti}] Row[{ri}] Cell[{ci}] keywords: {found}")
                print(f"  Text: {text[:150]}")

if not found_any:
    print("\n✅ No concurrency/anti-overselling content found!")

print("\n" + "=" * 80)
print("CHECK SPECIFIC MODIFIED PARAGRAPHS")
print("=" * 80)
for idx in [32, 38, 58, 78, 91, 96, 117, 118, 146, 179, 189, 190, 206, 222, 223, 225, 226, 228, 235]:
    p = doc.paragraphs[idx]
    text = p.text.strip()
    print(f"\nP[{idx}]: {text[:120]}...")
