import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

keywords = ['并发', '超卖', '分布式锁', 'Redisson', '限流', 'rateLimiter', 'RateLimiter', 'RedissonClient', 'getLock', 'tryLock', '高并发', '防超卖', '并发购买', '并发场景', '并发下', '500 并发', '500并发', 'JMeter', 'jmeter', '压测', '施压']

print("=" * 80)
print("SEARCHING FOR CONCURRENCY/ANTI-OVERSELLING RELATED CONTENT")
print("=" * 80)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    found = []
    for kw in keywords:
        if kw in text:
            found.append(kw)
    if found:
        print(f"\nP[{i}] keywords: {found}")
        print(f"  Text: {text[:150]}")
        if len(text) > 150:
            print(f"  ...{text[150:300]}")

print("\n" + "=" * 80)
print("SEARCHING IN TABLES")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if not text:
                continue
            found = []
            for kw in keywords:
                if kw in text:
                    found.append(kw)
            if found:
                print(f"\nTable[{ti}] Row[{ri}] Cell[{ci}] keywords: {found}")
                print(f"  Text: {text[:150]}")
