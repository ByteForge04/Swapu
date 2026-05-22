import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

print("=" * 70)
print("FINAL FORMAT VERIFICATION REPORT")
print("=" * 70)

print("\n--- 1. PAGE MARGINS ---")
expected = {'top': 2.5, 'bottom': 2.0, 'left': 2.5, 'right': 2.0, 'header': 2.6, 'footer': 2.4}
for i, sec in enumerate(doc.sections):
    ok_top = abs(sec.top_margin.cm - 2.5) < 0.01
    ok_bot = abs(sec.bottom_margin.cm - 2.0) < 0.01
    ok_left = abs(sec.left_margin.cm - 2.5) < 0.01
    ok_right = abs(sec.right_margin.cm - 2.0) < 0.01
    ok_hdr = abs(sec.header_distance.cm - 2.6) < 0.01
    ok_ftr = abs(sec.footer_distance.cm - 2.4) < 0.01
    status = "✅" if all([ok_top, ok_bot, ok_left, ok_right, ok_hdr, ok_ftr]) else "❌"
    print(f"  Section {i}: {status} top={sec.top_margin.cm:.2f} bot={sec.bottom_margin.cm:.2f} left={sec.left_margin.cm:.2f} right={sec.right_margin.cm:.2f} hdr={sec.header_distance.cm:.2f} ftr={sec.footer_distance.cm:.2f}")

print("\n--- 2. HEADER ---")
for i, sec in enumerate(doc.sections):
    header = sec.header
    for hp in header.paragraphs:
        if hp.text.strip():
            for run in hp.runs:
                size_pt = run.font.size / 12700 if run.font.size else None
                rPr = run._element.find(qn('w:rPr'))
                ea = ""
                if rPr is not None:
                    rF = rPr.find(qn('w:rFonts'))
                    if rF is not None:
                        ea = rF.get(qn('w:eastAsia'), '')
                ok = size_pt == 10.5 and ea == '宋体'
                status = "✅" if ok else "❌"
                print(f"  Section {i}: {status} text='{run.text}' font={ea} size={size_pt}pt")

print("\n--- 3. FOOTER (page number) ---")
for i, sec in enumerate(doc.sections):
    footer = sec.footer
    has_field = False
    for fp in footer.paragraphs:
        for child in fp._element:
            if child.tag == qn('w:r'):
                for sub in child:
                    if sub.tag == qn('w:instrText') and 'PAGE' in (sub.text or ''):
                        has_field = True
    sectPr = sec._sectPr
    pgNum = sectPr.find(qn('w:pgNumType'))
    fmt = pgNum.get(qn('w:fmt')) if pgNum is not None else None
    start = pgNum.get(qn('w:start')) if pgNum is not None else None
    print(f"  Section {i}: has_PAGE_field={has_field} fmt={fmt} start={start}")

print("\n--- 4. TITLE FORMATS ---")
checks = [
    (29, "摘要标题", "黑体", 18, None, WD_ALIGN_PARAGRAPH.CENTER),
    (33, "关键词行", None, None, None, None),
    (35, "Abstract标题", "Times New Roman", 18, True, WD_ALIGN_PARAGRAPH.CENTER),
    (39, "Key Words行", None, None, None, None),
    (41, "第1章 绪论", "黑体", 18, True, WD_ALIGN_PARAGRAPH.LEFT),
    (42, "1.1 节标题", "黑体", 16, True, WD_ALIGN_PARAGRAPH.LEFT),
    (170, "4.1.1 三级标题", "黑体", 14, True, WD_ALIGN_PARAGRAPH.LEFT),
    (241, "参考文献标题", "黑体", 18, True, WD_ALIGN_PARAGRAPH.LEFT),
    (262, "致谢标题", "黑体", 18, None, WD_ALIGN_PARAGRAPH.CENTER),
]

for idx, name, exp_font, exp_size, exp_bold, exp_align in checks:
    p = doc.paragraphs[idx]
    text = p.text.strip()[:30]
    r0 = p.runs[0] if p.runs else None
    if r0:
        rPr = r0._element.find(qn('w:rPr'))
        ea = ""
        if rPr is not None:
            rF = rPr.find(qn('w:rFonts'))
            if rF is not None:
                ea = rF.get(qn('w:eastAsia'), '')
        size_pt = r0.font.size / 12700 if r0.font.size else None
        bold = r0.font.bold
        align = p.paragraph_format.alignment
        
        issues = []
        if exp_font and exp_font not in (ea or ''):
            issues.append(f"font={ea}≠{exp_font}")
        if exp_size and abs(size_pt - exp_size) > 0.5:
            issues.append(f"size={size_pt}≠{exp_size}")
        if exp_bold is not None and bold != exp_bold:
            issues.append(f"bold={bold}≠{exp_bold}")
        if exp_align is not None and align != exp_align:
            issues.append(f"align={align}≠{exp_align}")
        
        status = "✅" if not issues else "❌ " + ", ".join(issues)
        print(f"  P[{idx}] {name}: {status} | '{text}' font={ea} size={size_pt}pt bold={bold} align={align}")
    else:
        print(f"  P[{idx}] {name}: ❌ no runs")

print("\n--- 5. BODY TEXT LINE SPACING ---")
body_indices = [30, 43, 47, 53, 71, 144, 168, 171, 205, 231]
for idx in body_indices:
    p = doc.paragraphs[idx]
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line = spacing.get(qn('w:line'))
            rule = spacing.get(qn('w:lineRule'))
            ok = line == '400' and rule == 'exact'
            status = "✅" if ok else "❌"
            print(f"  P[{idx}]: {status} line={line} rule={rule} (expected: 400/exact=20pt fixed)")

print("\n--- 6. REFERENCE SECTION ---")
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '参考文献':
        ref_idx = i
        break

if ref_idx:
    for j in range(ref_idx + 1, min(ref_idx + 4, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        if p.runs:
            r0 = p.runs[0]
            size_pt = r0.font.size / 12700 if r0.font.size else None
            rPr = r0._element.find(qn('w:rPr'))
            ea = ""
            if rPr is not None:
                rF = rPr.find(qn('w:rFonts'))
                if rF is not None:
                    ea = rF.get(qn('w:eastAsia'), '')
            ok = abs(size_pt - 10.5) < 0.5 and ea == '宋体'
            status = "✅" if ok else "❌"
            print(f"  P[{j}]: {status} font={ea} size={size_pt}pt | {p.text[:50]}")

print("\n--- 7. KEYWORDS DETAIL ---")
p33 = doc.paragraphs[33]
print(f"  关键词行 runs:")
for j, run in enumerate(p33.runs):
    rPr = run._element.find(qn('w:rPr'))
    ea = ""
    if rPr is not None:
        rF = rPr.find(qn('w:rFonts'))
        if rF is not None:
            ea = rF.get(qn('w:eastAsia'), '')
    size_pt = run.font.size / 12700 if run.font.size else None
    print(f"    Run[{j}]: '{run.text[:30]}' font={ea} size={size_pt}pt bold={run.font.bold}")

p39 = doc.paragraphs[39]
print(f"  Key Words行 runs:")
for j, run in enumerate(p39.runs):
    rPr = run._element.find(qn('w:rPr'))
    ea = ""
    if rPr is not None:
        rF = rPr.find(qn('w:rFonts'))
        if rF is not None:
            ea = rF.get(qn('w:eastAsia'), '')
    size_pt = run.font.size / 12700 if run.font.size else None
    print(f"    Run[{j}]: '{run.text[:30]}' font={ea} size={size_pt}pt bold={run.font.bold}")

print("\n" + "=" * 70)
print("VERIFICATION COMPLETE")
