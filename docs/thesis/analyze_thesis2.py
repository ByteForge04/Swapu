import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v11.docx')

print("=" * 60)
print("ALL SECTIONS - HEADER/FOOTER DETAILS")
print("=" * 60)
for i, sec in enumerate(doc.sections):
    print(f"\nSection {i}:")
    header = sec.header
    if header:
        print(f"  Header linked_to_previous: {header.is_linked_to_previous}")
        for j, hp in enumerate(header.paragraphs):
            if hp.text.strip() or hp.runs:
                print(f"  Header P[{j}]: '{hp.text}'")
                for run in hp.runs:
                    font = run.font
                    rPr = run._element.find(qn('w:rPr'))
                    ea_font = ""
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ea_font = rFonts.get(qn('w:eastAsia'), '')
                    print(f"    Run: '{run.text}' font={font.name}/{ea_font} size={font.size} bold={font.bold}")
    
    footer = sec.footer
    if footer:
        print(f"  Footer linked_to_previous: {footer.is_linked_to_previous}")
        for j, fp in enumerate(footer.paragraphs):
            if fp.text.strip() or fp.runs:
                print(f"  Footer P[{j}]: '{fp.text}'")
                for run in fp.runs:
                    font = run.font
                    print(f"    Run: '{run.text}' font={font.name} size={font.size}")
            for child in fp._element:
                if child.tag == qn('w:r'):
                    for sub in child:
                        if sub.tag in (qn('w:fldChar'), qn('w:instrText')):
                            val = sub.text if sub.text else sub.get(qn('w:fldCharType'), '')
                            print(f"    Field: {sub.tag.split('}')[1]} = {val}")

print("\n" + "=" * 60)
print("PARAGRAPHS 100-260 (remaining)")
print("=" * 60)
for i, p in enumerate(doc.paragraphs):
    if i < 100:
        continue
    if i >= 260:
        break
    text = p.text.strip()
    if not text and not p.runs:
        continue
    style_name = p.style.name if p.style else "None"
    pf = p.paragraph_format
    alignment = pf.alignment
    
    font_info = ""
    if p.runs:
        r = p.runs[0]
        font = r.font
        rPr = r._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        font_info = f"font={font.name}/{ea_font} size={font.size} bold={font.bold}"
    
    outline_lvl = ""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl = f" outlineLvl={outlineLvl.get(qn('w:val'))}"
    
    display_text = text[:60] if text else "(empty)"
    print(f"P[{i}] style='{style_name}' align={alignment} {font_info}{outline_lvl} | {display_text}")

print("\n" + "=" * 60)
print("TOTAL PARAGRAPHS:", len(doc.paragraphs))
print("TOTAL TABLES:", len(doc.tables))
print("TOTAL SECTIONS:", len(doc.sections))
