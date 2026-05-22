import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import copy

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v11.docx')

def set_run_font(run, cn_font, en_font, size_pt, bold=None):
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

def set_paragraph_line_spacing_fixed(paragraph, pt_value):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt_value * 20)))
    spacing.set(qn('w:lineRule'), 'exact')

def add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar1)
    
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE \\* MERGEFORMAT '
    run2._element.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar2)
    
    run4 = paragraph.add_run('1')
    set_run_font(run4, '宋体', 'Times New Roman', 9, bold=None)
    
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar3)

def set_header_text(section, text):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', 9, bold=None)

def set_footer_page_number(section, font_size_pt=9):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    add_page_number_field(fp)

print("Step 1: Fix page margins and header/footer distances...")
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(2.6)
    sec.footer_distance = Cm(2.4)

print("Step 2: Fix headers - '武汉理工大学本科毕业设计（论文）'...")
for i, sec in enumerate(doc.sections):
    set_header_text(sec, '武汉理工大学本科毕业设计（论文）')

print("Step 3: Fix footers with page numbers...")
for i, sec in enumerate(doc.sections):
    set_footer_page_number(sec)

print("Step 4: Fix page number format per section...")
sec0 = doc.sections[0]
sec1 = doc.sections[1]
sec2 = doc.sections[2]
sec3 = doc.sections[3]

sectPr0 = sec0._sectPr
pgNum0 = sectPr0.find(qn('w:pgNumType'))
if pgNum0 is None:
    pgNum0 = OxmlElement('w:pgNumType')
    sectPr0.append(pgNum0)
pgNum0.set(qn('w:fmt'), 'upperRoman')
pgNum0.set(qn('w:start'), '1')

sectPr1 = sec1._sectPr
pgNum1 = sectPr1.find(qn('w:pgNumType'))
if pgNum1 is None:
    pgNum1 = OxmlElement('w:pgNumType')
    sectPr1.append(pgNum1)
pgNum1.set(qn('w:fmt'), 'upperRoman')

sectPr2 = sec2._sectPr
pgNum2 = sectPr2.find(qn('w:pgNumType'))
if pgNum2 is not None:
    sectPr2.remove(pgNum2)
pgNum2 = OxmlElement('w:pgNumType')
sectPr2.append(pgNum2)
pgNum2.set(qn('w:fmt'), 'upperRoman')
pgNum2.set(qn('w:start'), '1')

sectPr3 = sec3._sectPr
pgNum3 = sectPr3.find(qn('w:pgNumType'))
if pgNum3 is not None:
    sectPr3.remove(pgNum3)
pgNum3 = OxmlElement('w:pgNumType')
sectPr3.append(pgNum3)
pgNum3.set(qn('w:fmt'), 'decimal')
pgNum3.set(qn('w:start'), '1')

print("Step 5: Fix paragraph fonts and line spacing...")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text and not p.runs:
        continue
    
    pPr = p._element.find(qn('w:pPr'))
    outline_lvl_val = None
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl_val = outlineLvl.get(qn('w:val'))
    
    is_chapter_title = False
    is_section_title = False
    is_subsection_title = False
    is_abstract_title = False
    is_ref_title = False
    is_thanks_title = False
    is_conclusion_title = False
    is_appendix_title = False
    is_keyword_line = False
    is_fig_caption = False
    is_table_caption = False
    is_toc_entry = False
    
    if outline_lvl_val == '0':
        if text.startswith('第') and '章' in text[:5]:
            is_chapter_title = True
        elif text == '参考文献':
            is_ref_title = True
        elif text == '致谢' or text == '致 谢':
            is_thanks_title = True
        elif '结论' in text and text.startswith('第'):
            is_conclusion_title = True
            is_chapter_title = True
        elif '总结' in text and text.startswith('第'):
            is_chapter_title = True
        else:
            is_chapter_title = True
    elif outline_lvl_val == '1':
        is_section_title = True
    elif outline_lvl_val == '2':
        is_subsection_title = True
    
    if not is_chapter_title and not is_section_title and not is_subsection_title:
        if text == '摘要' or text == '摘    要':
            is_abstract_title = True
        elif text == 'Abstract':
            is_abstract_title = True
        elif text.startswith('关键词') or text.startswith('Key Words') or text.startswith('Keywords'):
            is_keyword_line = True
        elif text.startswith('图 ') or text.startswith('图2-') or text.startswith('图3-') or text.startswith('图4-'):
            is_fig_caption = True
        elif text.startswith('表 ') or text.startswith('表5-') or text.startswith('表2-') or text.startswith('表3-'):
            is_table_caption = True
    
    style_name = p.style.name if p.style else ""
    if 'toc' in style_name.lower():
        is_toc_entry = True
    
    if i <= 25:
        continue
    
    if is_chapter_title or is_ref_title or is_thanks_title or is_conclusion_title:
        for run in p.runs:
            set_run_font(run, '黑体', 'Times New Roman', 18, bold=True)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr2 = p._element.get_or_add_pPr()
        spacing = pPr2.find(qn('w:spacing'))
        if spacing is not None:
            spacing.attrib.pop(qn('w:line'), None)
            spacing.attrib.pop(qn('w:lineRule'), None)
    
    elif is_abstract_title:
        for run in p.runs:
            if text == 'Abstract':
                set_run_font(run, 'Times New Roman', 'Times New Roman', 18, bold=True)
            else:
                set_run_font(run, '黑体', 'Times New Roman', 18, bold=None)
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    elif is_section_title:
        for run in p.runs:
            set_run_font(run, '黑体', 'Times New Roman', 16, bold=True)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr2 = p._element.get_or_add_pPr()
        spacing = pPr2.find(qn('w:spacing'))
        if spacing is not None:
            spacing.attrib.pop(qn('w:line'), None)
            spacing.attrib.pop(qn('w:lineRule'), None)
    
    elif is_subsection_title:
        for run in p.runs:
            set_run_font(run, '黑体', 'Times New Roman', 14, bold=True)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr2 = p._element.get_or_add_pPr()
        spacing = pPr2.find(qn('w:spacing'))
        if spacing is not None:
            spacing.attrib.pop(qn('w:line'), None)
            spacing.attrib.pop(qn('w:lineRule'), None)
    
    elif is_keyword_line:
        if text.startswith('关键词'):
            for run in p.runs:
                run_text = run.text
                if '关键词' in run_text or '：' in run_text:
                    set_run_font(run, '黑体', 'Times New Roman', 14, bold=None)
                else:
                    set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
        elif text.startswith('Key Words') or text.startswith('Keywords'):
            for run in p.runs:
                run_text = run.text
                if 'Key Words' in run_text or 'Keywords' in run_text or '：' in run_text or ':' in run_text:
                    set_run_font(run, 'Times New Roman', 'Times New Roman', 14, bold=True)
                else:
                    set_run_font(run, 'Times New Roman', 'Times New Roman', 12, bold=None)
    
    elif is_fig_caption or is_table_caption:
        for run in p.runs:
            set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    elif is_ref_title:
        pass
    
    elif is_toc_entry:
        pass
    
    else:
        is_body_text = True
        if i >= 241 and text.startswith('['):
            is_body_text = False
        
        if is_body_text:
            for run in p.runs:
                rPr = run._element.find(qn('w:rPr'))
                vertAlign = None
                if rPr is not None:
                    vertAlign = rPr.find(qn('w:vertAlign'))
                
                if vertAlign is not None and vertAlign.get(qn('w:val')) == 'superscript':
                    set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
                else:
                    set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
            
            set_paragraph_line_spacing_fixed(p, 20)

print("Step 6: Fix reference section font size (宋体五号=10.5pt)...")
ref_started = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == '参考文献':
        ref_started = True
        continue
    if ref_started:
        if text.startswith('['):
            for run in p.runs:
                set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=None)
            set_paragraph_line_spacing_fixed(p, 20)
        elif text == '致谢' or text == '致 谢':
            ref_started = False

print("Step 7: Fix 致谢 section...")
thanks_started = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == '致谢' or text == '致 谢':
        thanks_started = True
        for run in p.runs:
            set_run_font(run, '黑体', 'Times New Roman', 18, bold=None)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue
    if thanks_started:
        if text:
            for run in p.runs:
                set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
            set_paragraph_line_spacing_fixed(p, 20)

print("Step 8: Fix abstract body text...")
abstract_started = False
abstract_ended = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == '摘要' or text == '摘    要':
        abstract_started = True
        continue
    if abstract_started and not abstract_ended:
        if text.startswith('关键词'):
            abstract_ended = True
            continue
        if text:
            for run in p.runs:
                set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
            set_paragraph_line_spacing_fixed(p, 20)

en_abstract_started = False
en_abstract_ended = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == 'Abstract':
        en_abstract_started = True
        continue
    if en_abstract_started and not en_abstract_ended:
        if text.startswith('Key Words') or text.startswith('Keywords'):
            en_abstract_ended = True
            continue
        if text:
            for run in p.runs:
                set_run_font(run, 'Times New Roman', 'Times New Roman', 12, bold=None)
            set_paragraph_line_spacing_fixed(p, 20)

print("Step 9: Fix table cell fonts...")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    vertAlign = None
                    if rPr is not None:
                        vertAlign = rPr.find(qn('w:vertAlign'))
                    if vertAlign is not None and vertAlign.get(qn('w:val')) == 'superscript':
                        set_run_font(run, '宋体', 'Times New Roman', 12, bold=None)
                    else:
                        current_size = run.font.size
                        if current_size and current_size < Pt(9):
                            set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=None)

output_path = r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
