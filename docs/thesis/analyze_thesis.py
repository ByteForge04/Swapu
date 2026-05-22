import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v11.docx')

print("=" * 60)
print("PAGE SETUP")
print("=" * 60)
for i, sec in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  top_margin: {sec.top_margin.cm:.2f} cm")
    print(f"  bottom_margin: {sec.bottom_margin.cm:.2f} cm")
    print(f"  left_margin: {sec.left_margin.cm:.2f} cm")
    print(f"  right_margin: {sec.right_margin.cm:.2f} cm")
    print(f"  header_distance: {sec.header_distance.cm:.2f} cm")
    print(f"  footer_distance: {sec.footer_distance.cm:.2f} cm")
    print(f"  gutter: {sec.gutter.cm:.2f} cm")
    print(f"  different_first_page_header_footer: {sec.different_first_page_header_footer}")
    
    header = sec.header
    if header and header.paragraphs:
        for hp in header.paragraphs:
            if hp.text.strip():
                print(f"  Header text: '{hp.text}'")
                for run in hp.runs:
                    font = run.font
                    rPr = run._element.find(qn('w:rPr'))
                    ea_font = ""
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ea_font = rFonts.get(qn('w:eastAsia'), '')
                    print(f"    Run: '{run.text}' font={font.name}/{ea_font} size={font.size} bold={font.bold}")
    
    footer = sec.footer
    if footer and footer.paragraphs:
        for fp in footer.paragraphs:
            if fp.text.strip():
                print(f"  Footer text: '{fp.text}'")
            for child in fp._element:
                if child.tag == qn('w:r'):
                    for sub in child:
                        if sub.tag == qn('w:fldChar') or sub.tag == qn('w:instrText'):
                            print(f"  Footer field: {sub.tag} = {sub.text if sub.text else sub.get(qn('w:fldCharType'), '')}")

print("\n" + "=" * 60)
print("PARAGRAPHS (first 100)")
print("=" * 60)
for i, p in enumerate(doc.paragraphs):
    if i >= 100:
        break
    text = p.text.strip()
    if not text and not p.runs:
        continue
    style_name = p.style.name if p.style else "None"
    pf = p.paragraph_format
    alignment = pf.alignment
    
    font_info = ""
    if p.runs:
        r = p.runs[0]
        font = r.font
        rPr = r._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        font_info = f"font={font.name}/{ea_font} size={font.size} bold={font.bold}"
    
    outline_lvl = ""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl = f" outlineLvl={outlineLvl.get(qn('w:val'))}"
    
    display_text = text[:60] if text else "(empty)"
    print(f"P[{i}] style='{style_name}' align={alignment} {font_info}{outline_lvl} | {display_text}")
