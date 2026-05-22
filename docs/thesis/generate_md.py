import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
import os

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v14.docx')

image_map = {}
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image = rel.target_part
        image_bytes = image.blob
        import hashlib
        h = hashlib.md5(image_bytes).hexdigest()[:8]
        ext = os.path.splitext(image.partname)[1] or '.png'
        image_map[rel.rId] = f"thesis_images/img_{h}{ext}"

para_image_map = {}
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        drawing_elements = run._element.findall('.//' + qn('wp:inline'))
        for shape in drawing_elements:
            blip = shape.find('.//' + qn('a:blip'))
            if blip is not None:
                rId = blip.get(qn('r:embed'))
                if rId in image_map:
                    if i not in para_image_map:
                        para_image_map[i] = []
                    para_image_map[i].append(image_map[rId])

md_lines = []
md_lines.append("# 基于RAG的校园智能二手交易系统设计与实现\n")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    pPr = p._element.find(qn('w:pPr'))
    outline_lvl = None
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl = outlineLvl.get(qn('w:val'))
    
    if i in para_image_map:
        for img_path in para_image_map[i]:
            md_lines.append(f"\n![图片]({img_path})\n")
    
    if not text:
        continue
    
    if i <= 27:
        if i == 0:
            continue
        if '原创性声明' in text or '使用授权书' in text:
            md_lines.append(f"## {text}\n")
        elif text == '摘    要':
            md_lines.append(f"## 摘  要\n")
        elif text == 'Abstract':
            md_lines.append(f"## Abstract\n")
        else:
            md_lines.append(f"{text}\n")
        continue
    
    if outline_lvl == '0':
        md_lines.append(f"\n## {text}\n")
    elif outline_lvl == '1':
        md_lines.append(f"\n### {text}\n")
    elif outline_lvl == '2':
        md_lines.append(f"\n#### {text}\n")
    elif text.startswith('图 ') or text.startswith('图2-') or text.startswith('图3-') or text.startswith('图4-') or text.startswith('图5-'):
        md_lines.append(f"\n*{text}*\n")
    elif text.startswith('表 ') or text.startswith('表5-') or text.startswith('表2-'):
        md_lines.append(f"\n**{text}**\n")
    elif text.startswith('[') and ']' in text[:5] and i > 240:
        md_lines.append(f"{text}\n")
    else:
        md_lines.append(f"{text}\n")

output_path = r'd:\SwapU\docs\thesis\SwapU毕业论文.md'
with open(output_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(md_lines))

print(f"Written to {output_path}")
print(f"Total lines: {len(md_lines)}")
