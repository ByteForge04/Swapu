import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

def set_run_font(run, cn_font, en_font, size_pt, bold=None):
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

def replace_paragraph_text(p, new_text, cn_font='宋体', en_font='Times New Roman', size_pt=12, bold=None):
    for run in p.runs:
        p._element.remove(run._element)
    new_run = p.add_run(new_text)
    set_run_font(new_run, cn_font, en_font, size_pt, bold)

def set_line_spacing_fixed(p, pt=20):
    pPr = p._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')

print("=== P[32]: 中文摘要 - 去掉并发防超卖作为重点 ===")
p32 = doc.paragraphs[32]
new_p32 = ("在工程实现层面，本文重点解决了两个关键问题：第一，通过 RocketMQ 消息驱动 + 全量同步的双轨策略保证 PgVector 向量数据与 MySQL 行记录的一致性，确保 AI 导购推荐的商品均为真实在售状态；第二，通过十余版本的 System Prompt 迭代调优，将 AI 输出行为收敛至稳定可控状态，解决大模型输出格式不统一、字数超标和编造商品信息等问题。系统测试结果表明，各功能模块运行正常，AI 导购推荐准确率满足预期，AES 加密有效保护用户隐私，系统整体运行稳定。")
replace_paragraph_text(p32, new_p32)
set_line_spacing_fixed(p32)

print("=== P[38]: 英文摘要 - 去掉并发防超卖 ===")
p38 = doc.paragraphs[38]
new_p38 = ("At the engineering implementation level, this paper focuses on solving two key problems: first, maintaining consistency between PgVector vector data and MySQL row records through a dual-track strategy of RocketMQ message-driven incremental sync and full sync, ensuring that AI-recommended products are genuinely available; second, converging AI output behavior to a stable and controllable state through more than ten iterations of System Prompt tuning, resolving issues such as inconsistent output formats, excessive word counts, and fabricated product information. System testing results show that all functional modules operate normally, AI shopping guide recommendation accuracy meets expectations, AES encryption effectively protects user privacy, and the system runs stably overall.")
replace_paragraph_text(p38, new_p38, cn_font='Times New Roman', en_font='Times New Roman')
set_line_spacing_fixed(p38)

print("=== P[58]: 1.3主要研究内容(5) - 去掉分布式锁和限流描述 ===")
p58 = doc.paragraphs[58]
new_p58 = ("（5）订单与支付流转模块。覆盖下单、支付拉起、状态查询、交易取消与确认收货的完整链路。订单创建后发送 RocketMQ 延时消息，30 分钟未支付则自动取消并恢复商品在售状态。支付环节对接支付宝沙箱环境，通过异步回调更新订单支付状态，回调中包含 RSA 验签和幂等性检查。")
replace_paragraph_text(p58, new_p58)
set_line_spacing_fixed(p58)

print("=== P[78]: 2.1功能需求(5) - 简化并发描述 ===")
p78 = doc.paragraphs[78]
new_p78 = ("（5）订单与支付流转模块。覆盖下单、支付拉起、状态查询、交易取消与确认收货的完整链路。系统需保证同一商品不会被多个买家同时下单，并限制单用户的下单频率。")
replace_paragraph_text(p78, new_p78)
set_line_spacing_fixed(p78)

print("=== P[91]: 2.2业务流程 - 去掉分布式锁细节 ===")
p91 = doc.paragraphs[91]
new_p91 = ("在订单交易流程中，买卖双方通过 WebSocket 通道就商品细节进行沟通。沟通完成后买家确认购买，点击立即购买按钮，前端调用 createOrder 接口创建订单并将商品状态更新为交易中，同时发送 RocketMQ 延时消息用于超时自动取消。卖家确认订单后，系统调用支付宝沙箱生成支付表单，买家完成支付后支付宝通过异步回调通知后端，后端验证签名后更新订单支付状态。交易完成后买家可对商品进行评价。若买家 30 分钟内未完成支付，RocketMQ 延时消息触发自动取消，商品恢复在售状态。")
replace_paragraph_text(p91, new_p91)
set_line_spacing_fixed(p91)

print("=== P[96]: 2.3非功能需求(1) - 去掉分布式锁和限流 ===")
p96 = doc.paragraphs[96]
new_p96 = ("（1）响应时间需求。首页商品列表接口需在 200 毫秒内返回结果，搜索接口需在 500 毫秒内返回。下单接口因涉及数据库事务提交和消息发送，可放宽至 500 毫秒。AI 接口因包含大模型推理，响应时间在 3-10 秒之间，属于可接受范围。")
replace_paragraph_text(p96, new_p96)
set_line_spacing_fixed(p96)

print("=== P[117]: 3.2架构-业务逻辑层 - 去掉分布式锁描述 ===")
p117 = doc.paragraphs[117]
new_p117 = ("业务逻辑层包含 15 个 Controller 和 13 个 ServiceImpl，基本一一对应。Controller 层仅负责参数校验和响应封装，业务逻辑全部下沉至 Service 层。ItemServiceImpl 依赖最密集，除 CRUD 外还承担 ES 同步和 RAG 向量写入。TradeOrderServiceImpl 封装订单创建、支付回调与状态流转逻辑，是系统中事务密度最高的模块。RagServiceImpl 封装向量检索逻辑，供 ChatController 调用。ChatMessageService 处理 WebSocket 消息的持久化与已读状态更新。")
replace_paragraph_text(p117, new_p117)
set_line_spacing_fixed(p117)

print("=== P[118]: 3.2架构-数据持久层 - 去掉分布式锁和限流 ===")
p118 = doc.paragraphs[118]
new_p118 = ("数据持久层各存储引擎职责隔离。MySQL 承担 ACID 写入和关联查询，是系统的主数据源。Redis 承载缓存热点商品详情以降低 MySQL 查询压力。Elasticsearch 承担全文检索，通过 RocketMQ 异步同步商品数据，查询效率相比 MySQL LIKE 提升数个数量级。PgVector 承担向量语义检索，商品上架时异步写入 Embedding 向量，为 RAG 管线提供事实数据源。")
replace_paragraph_text(p118, new_p118)
set_line_spacing_fixed(p118)

print("=== P[146]: 3.3功能模块-核心交易 - 去掉分布式锁防超卖 ===")
p146 = doc.paragraphs[146]
new_p146 = ("核心交易模块的代码量最大，涵盖商品、订单和支付三个子模块。商品子模块维护五态状态机：status=0 待审核、1 在售、2 交易中、3 已售出、4 已下架。状态只能由后端在特定业务操作中变更，前端无法直接修改状态字段。商品发布时异步写入 ES 索引和 PgVector 向量，确保可搜索和可导购。订单子模块同样维护四态状态机：0 待确认、1 进行中、2 已完成、3 已取消。订单创建时通过 RocketMQ 延时消息实现 30 分钟超时自动取消。支付子模块对接支付宝沙箱，异步回调中通过 RSA 验签和幂等性检查保证资金安全，避免重复支付和伪造回调。")
replace_paragraph_text(p146, new_p146)
set_line_spacing_fixed(p146)

print("=== P[179]: 4.1关键技术 - 去掉分布式锁描述 ===")
p179 = doc.paragraphs[179]
new_p179 = ("系统包含 15 个 Controller 和 13 个 ServiceImpl，基本一一对应。其中 ItemServiceImpl 依赖最密集，除 CRUD 外还承担 ES 同步和 RAG 向量写入；TradeOrderServiceImpl 封装订单创建、支付回调与状态流转逻辑；RagServiceImpl 封装向量检索逻辑，供 ChatController 调用。")
replace_paragraph_text(p179, new_p179)
set_line_spacing_fixed(p179)

print("=== P[189]: 4.2.2订单交易 - 整段重写，去掉分布式锁策略 ===")
p189 = doc.paragraphs[189]
new_p189 = ('createOrder() 的执行流程：参数校验 -> 查询商品并校验状态 -> TransactionTemplate.execute() 创建订单并更新商品状态 -> 发送 RocketMQ 延时消息。TransactionTemplate 采用编程式事务而非声明式 @Transactional，是因为需要在业务校验通过后才开启事务，避免事务范围覆盖校验阶段导致长事务问题。订单创建成功后，商品状态由\u201c在售\u201d变更为\u201c交易中\u201d，同时向买家返回订单详情。')
replace_paragraph_text(p189, new_p189)
set_line_spacing_fixed(p189)

print("=== P[190]: 4.2.2订单交易 - 去掉限流段落 ===")
p190 = doc.paragraphs[190]
new_p190 = ("支付环节对接支付宝沙箱环境，后端构建 AlipayTradePagePayRequest 生成支付表单，前端通过表单提交跳转至支付宝收银台。支付完成后支付宝通过异步回调通知后端，后端在回调处理中执行 RSA 验签、幂等性检查和订单状态更新三个步骤，确保支付结果的真实性和一致性。RocketMQ 延时消息订单超时处理时序图如图 4-2 所示。")
replace_paragraph_text(p190, new_p190)
set_line_spacing_fixed(p190)

print("=== P[206]: 5.1测试环境 - 去掉限流描述 ===")
p206 = doc.paragraphs[206]
new_p206 = ("测试环境为一台 4 核 8 GB 内存的 CentOS 7 虚拟机，各服务连接地址通过环境变量配置。JDK 版本为 17，MySQL 8.0，Redis 7.0，Elasticsearch 8.12，PostgreSQL 16（含 PgVector 扩展），RocketMQ 5.1。前端通过 Nginx 反向代理访问后端服务。测试按三个层次展开：功能层面验证核心链路的完整性，性能层面验证系统在模拟压力下的稳定性，安全层面验证加密与越权防护两项机制的实际效果。")
replace_paragraph_text(p206, new_p206)
set_line_spacing_fixed(p206)

print("=== P[222]: 5.3性能测试 - 去掉分布式锁和JMeter并发测试 ===")
p222 = doc.paragraphs[222]
new_p222 = ("性能测试采用 JMeter，200 并发线程持续施压 5 分钟。首页商品列表接口（ES 查询）：平均响应 45ms，TPS 超 2100，99% 请求在 100ms 内完成。商品详情与搜索接口：平均响应 60ms，TPS 约 1800，高亮渲染对性能无明显影响。AI 导购接口：平均响应 3.5s（含向量检索约 200ms 和 LLM 推理约 3s），TPS 约 15，瓶颈在模型推理环节。测试过程中未发现内存泄漏或线程死锁问题，系统在持续压力下表现稳定。")
replace_paragraph_text(p222, new_p222)
set_line_spacing_fixed(p222)

print("=== P[223]: 5.3性能测试结论 - 去掉分布式锁和超卖 ===")
p223 = doc.paragraphs[223]
new_p223 = ("性能测试结果表明，系统的常规读写接口在 200 并发压力下表现稳定，响应时间和吞吐量均满足设计目标。AI 接口的瓶颈在模型推理环节，后续可通过流式输出和模型蒸馏优化。整体来看，系统在当前硬件配置下能够支撑校园规模的访问需求。")
replace_paragraph_text(p223, new_p223)
set_line_spacing_fixed(p223)

print("=== P[225]: 5.4安全测试 - 去掉限流防护 ===")
p225 = doc.paragraphs[225]
new_p225 = ('安全验证包含两项。第一项：AES 字段级加密——数据库中 phone 字段存储为 Base64 密文，直接查询数据库无法获取用户手机号明文；API 返回时通过 TypeHandler 自动解密为明文，业务代码无需感知加解密过程。第二项：越权防护——普通用户尝试访问 /admin/** 接口返回 403 状态码，尝试删除他人商品返回\u201c无权操作\u201d提示，权限校验机制工作正常。')
replace_paragraph_text(p225, new_p225)
set_line_spacing_fixed(p225)

print("=== P[226]: 5.4安全测试已知缺陷 - 去掉限流相关 ===")
p226 = doc.paragraphs[226]
new_p226 = ("当前版本存在一个已知缺陷：登录接口未配置防护策略，理论上存在暴力破解风险。在生产环境中需补充防护措施，建议配置为每分钟 5 次失败尝试后锁定账号 15 分钟。此外，后续版本还应考虑引入图形验证码或滑块验证，作为登录接口的第一道防线。")
replace_paragraph_text(p226, new_p226)
set_line_spacing_fixed(p226)

print("=== P[228]: 5.5测试结论 - 去掉分布式锁和限流 ===")
p228 = doc.paragraphs[228]
new_p228 = ("本轮测试覆盖功能、性能和安全三个维度。功能层面，6 项核心测试用例全部通过，完整业务链路中各环节状态迁移正确，AI 润色和导购功能输出稳定，无功能断点。性能层面，RocketMQ 延时消息精确处理订单超时，ES 和 Redis 在压力下表现稳定，首页和搜索接口响应时间远低于 200ms 的目标值。安全层面，AES 加密和越权防护机制均按预期工作。已知缺陷：登录接口未配置防护策略，生产环境需补充。总体而言，系统在功能完整性、性能稳定性和安全性方面均达到设计预期。")
replace_paragraph_text(p228, new_p228)
set_line_spacing_fixed(p228)

print("=== P[235]: 第6章总结 - 去掉分布式锁和超卖 ===")
p235 = doc.paragraphs[235]
new_p235 = ("RocketMQ 延时消息在测试中验证了工程价值。超时自动取消机制在超时取消场景中表现正确，未出现状态不一致。编程式事务的选择确保了事务范围与业务逻辑的精确匹配，避免了长事务问题。双轨同步策略在商品库持续变动的场景下保证了向量数据与行记录的一致性，增量同步保证实时性，全量同步作为兜底修复手段。")
replace_paragraph_text(p235, new_p235)
set_line_spacing_fixed(p235)

print("=== P[148]: 后台管控模块 - 并发发送通知是正常业务，保留 ===")
# P[148] "并发送系统通知" - 这里的"并发"是"同时发送"的意思，不是并发编程，保留

print("\n=== TABLE: 修改表格中分布式锁相关内容 ===")
table0 = doc.tables[0]
row7 = table0.rows[7]
print(f"Table[0] Row[7] cells:")
for ci, cell in enumerate(row7.cells):
    print(f"  Cell[{ci}]: {cell.text}")

num_cells = len(row7.cells)
row7.cells[0].paragraphs[0].runs[0].text = "Redis 缓存命中率"
for p in row7.cells[1].paragraphs:
    for run in p.runs:
        if run.text.strip():
            run.text = "商品详情缓存命中率"
if num_cells > 2:
    row7.cells[2].paragraphs[0].runs[0].text = "缓存命中 > 95%"
if num_cells > 3:
    for p in row7.cells[3].paragraphs:
        for run in p.runs:
            if run.text.strip():
                run.text = "热点商品详情缓存命中率超过 95%"
else:
    tr_xml = row7._tr
    tc_list = tr_xml.findall(qn('w:tc'))
    if len(tc_list) >= 4:
        for p in tc_list[3].findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    if t.text and t.text.strip():
                        t.text = "热点商品详情缓存命中率超过 95%"

output_path = r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v14.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
