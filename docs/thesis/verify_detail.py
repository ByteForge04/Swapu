import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

print("=== Key Words line runs ===")
for idx in [33, 39]:
    p = doc.paragraphs[idx]
    print(f"P[{idx}] text: {p.text[:80]}")
    for j, run in enumerate(p.runs):
        rPr = run._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        print(f"  Run[{j}]: '{run.text}' font={run.font.name}/{ea_font} size={run.font.size} bold={run.font.bold}")

print("\n=== Abstract title runs ===")
for idx in [29, 35]:
    p = doc.paragraphs[idx]
    print(f"P[{idx}] text: {p.text}")
    for j, run in enumerate(p.runs):
        rPr = run._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        print(f"  Run[{j}]: '{run.text}' font={run.font.name}/{ea_font} size={run.font.size} bold={run.font.bold}")

print("\n=== Header font size check ===")
sec = doc.sections[0]
header = sec.header
for hp in header.paragraphs:
    for run in hp.runs:
        print(f"Header run: '{run.text}' size={run.font.size} = {run.font.size / 12700 if run.font.size else None}pt")

print("\n=== Page number font size check ===")
footer = sec.footer
for fp in footer.paragraphs:
    for run in fp.runs:
        if run.text:
            print(f"Footer run: '{run.text}' size={run.font.size} = {run.font.size / 12700 if run.font.size else None}pt")
