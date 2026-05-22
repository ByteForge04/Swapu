import sys
sys.path.insert(0, r'D:\SwapU\docs\thesis\pip_libs')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

doc = Document(r'd:\SwapU\docs\thesis\SwapU毕业论文_终版v13.docx')

print("=" * 60)
print("PAGE SETUP VERIFICATION")
print("=" * 60)
for i, sec in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  top_margin: {sec.top_margin.cm:.2f} cm (expected 2.50)")
    print(f"  bottom_margin: {sec.bottom_margin.cm:.2f} cm (expected 2.00)")
    print(f"  left_margin: {sec.left_margin.cm:.2f} cm (expected 2.50)")
    print(f"  right_margin: {sec.right_margin.cm:.2f} cm (expected 2.00)")
    print(f"  header_distance: {sec.header_distance.cm:.2f} cm (expected 2.60)")
    print(f"  footer_distance: {sec.footer_distance.cm:.2f} cm (expected 2.40)")
    
    header = sec.header
    if header and header.paragraphs:
        for hp in header.paragraphs:
            if hp.text.strip():
                print(f"  Header: '{hp.text}'")
                for run in hp.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    ea_font = ""
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ea_font = rFonts.get(qn('w:eastAsia'), '')
                    print(f"    font={run.font.name}/{ea_font} size={run.font.size} bold={run.font.bold}")
    
    footer = sec.footer
    if footer and footer.paragraphs:
        for fp in footer.paragraphs:
            has_field = False
            for child in fp._element:
                if child.tag == qn('w:r'):
                    for sub in child:
                        if sub.tag in (qn('w:fldChar'), qn('w:instrText')):
                            has_field = True
            if has_field or fp.text.strip():
                print(f"  Footer: has_page_number_field={has_field} text='{fp.text}'")
    
    sectPr = sec._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        print(f"  pgNumType: fmt={pgNumType.get(qn('w:fmt'))} start={pgNumType.get(qn('w:start'))}")

print("\n" + "=" * 60)
print("KEY PARAGRAPHS VERIFICATION")
print("=" * 60)

check_indices = [29, 30, 33, 35, 39, 41, 42, 46, 52, 60, 69, 70, 143, 154, 167, 170, 173, 203, 230, 241, 242, 262, 263]
for i in check_indices:
    if i >= len(doc.paragraphs):
        continue
    p = doc.paragraphs[i]
    text = p.text.strip()
    if not text:
        continue
    
    font_info = ""
    if p.runs:
        r = p.runs[0]
        rPr = r._element.find(qn('w:rPr'))
        ea_font = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'), '')
        font_info = f"font={r.font.name}/{ea_font} size={r.font.size} bold={r.font.bold}"
    
    outline_lvl = ""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            outline_lvl = f" lvl={outlineLvl.get(qn('w:val'))}"
    
    pf = p.paragraph_format
    ls_info = f" ls={pf.line_spacing}"
    
    display = text[:50] if text else "(empty)"
    print(f"P[{i}] {font_info}{outline_lvl}{ls_info} | {display}")

print("\n" + "=" * 60)
print("LINE SPACING CHECK")
print("=" * 60)
for i in [30, 43, 47, 52, 69, 167, 203, 230, 242, 263]:
    if i >= len(doc.paragraphs):
        continue
    p = doc.paragraphs[i]
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line = spacing.get(qn('w:line'))
            rule = spacing.get(qn('w:lineRule'))
            sb = spacing.get(qn('w:before'))
            sa = spacing.get(qn('w:after'))
            print(f"P[{i}] line={line} rule={rule} before={sb} after={sa} | {p.text[:40]}")
        else:
            print(f"P[{i}] no spacing element | {p.text[:40]}")
