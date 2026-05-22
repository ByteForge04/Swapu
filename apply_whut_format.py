import docx
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import re

def set_font(run, ascii_font='Times New Roman', east_asia_font='宋体', size=12, bold=False):
    run.font.name = ascii_font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
    run.font.size = Pt(size)
    run.font.bold = bold

def apply_whut_format(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 1. 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(2.6)
        section.footer_distance = Cm(2.4)
        
    # 状态
    in_cover = True
    in_ref = False
    in_ack = False
    
    # 我们需要确保某些标题前有分页符。
    # 避免在文档最开头插入分页符，所以记录是否是第一页
    first_page = True
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # 独创性声明 标志着封面结束
        if text == "独创性声明":
            in_cover = False
            
        if text == "参考文献":
            in_ref = True
        if text == "致谢":
            in_ref = False
            in_ack = True
            
        # --- 清理可能的旧分页符（如果存在于文本中，我们不一定能清掉，但尽量标准化） ---
        
        # 标题匹配逻辑
        is_chapter_title = bool(re.match(r'^第[一二三四五六七八九十0-9]+章\s+', text))
        is_major_title = text in ["摘  要", "摘要", "Abstract", "目  录", "目录", "参考文献", "致谢", "独创性声明", "学位论文使用授权书"]
        
        # --- 强制分页与标题格式 ---
        if is_chapter_title or is_major_title:
            # 确保标题在新的一页 (除了封面内部的)
            if not in_cover and not first_page:
                # 检查该段落前面是否有分页符
                has_page_break = False
                for run in para.runs:
                    if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        has_page_break = True
                        break
                # 如果前一个段落是空段落且包含分页符，也算
                if i > 0 and not has_page_break:
                    prev_p = doc.paragraphs[i-1]
                    for run in prev_p.runs:
                        if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                            has_page_break = True
                            break
                
                if not has_page_break:
                    # 在段首插入分页符
                    para.insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)
            
            if text in ["独创性声明", "学位论文使用授权书"]:
                first_page = False
                
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            
            # 段前段后0.5行 (约 0.5 * 1.5 * 12pt = 9pt)
            para.paragraph_format.space_before = Pt(9)
            para.paragraph_format.space_after = Pt(9)
            
            font_name = 'Times New Roman' if text == "Abstract" else '黑体'
            for run in para.runs: 
                set_font(run, ascii_font=font_name, east_asia_font=font_name, size=18, bold=True) # 小二号
            continue
            
        # 封面特殊处理
        if in_cover:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            if "武汉理工大学毕业设计（论文）" in text:
                for run in para.runs: set_font(run, east_asia_font='华文中宋', size=26, bold=True) # 一号
            elif text.startswith("基于"):
                for run in para.runs: set_font(run, east_asia_font='黑体', size=22, bold=True) # 二号
            else:
                for run in para.runs: set_font(run, east_asia_font='华文中宋', size=16) # 三号
            continue
            
        # --- 节标题处理 ---
        # 1.1 一级节标题
        if re.match(r'^\d+\.\d+\s+', text):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(9)
            para.paragraph_format.space_after = Pt(9)
            for run in para.runs: set_font(run, east_asia_font='黑体', size=16, bold=True) # 三号
            continue
            
        # 1.1.1 二级节标题
        if re.match(r'^\d+\.\d+\.\d+\s+', text):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs: set_font(run, east_asia_font='黑体', size=14, bold=True) # 四号
            continue
            
        # --- 图表标题 ---
        if re.match(r'^(图|表)\s*\d+-\d+', text) or text.startswith("代码清单"):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs: set_font(run, east_asia_font='宋体', size=12) # 小四号
            continue
            
        # --- 参考文献条目 ---
        if in_ref and re.match(r'^\[\d+\]', text):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = Pt(20) # 固定值20磅
            for run in para.runs: set_font(run, east_asia_font='宋体', size=10.5) # 五号
            continue
            
        # --- 正文段落 ---
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.line_spacing = Pt(20) # 行距：固定值20磅
        
        # 特殊处理：关键词
        if text.startswith("关键词：") or text.startswith("Key Words:"):
            para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs:
                if "关键词" in run.text or "Key Words" in run.text:
                    set_font(run, ascii_font='Times New Roman', east_asia_font='黑体', size=14, bold=True) # 黑体四号
                else:
                    set_font(run, ascii_font='Times New Roman', east_asia_font='宋体', size=12)
            continue
            
        # 代码段落
        if text.startswith("//") or text.startswith("rateLimiter") or text.startswith("RLock") or text.startswith("boolean") or "throw new" in text:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.left_indent = Pt(24)
            for run in para.runs: set_font(run, ascii_font='Courier New', east_asia_font='宋体', size=10)
            continue

        # 普通正文
        para.paragraph_format.first_line_indent = Pt(24) # 首行缩进2字符 (24磅)
        for run in para.runs:
            set_font(run, ascii_font='Times New Roman', east_asia_font='宋体', size=12)

    # 处理表格字体
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    for r in p.runs:
                        set_font(r, ascii_font='Times New Roman', east_asia_font='宋体', size=10.5) # 表格内通常用五号

    doc.save(output_path)
    print("WHUT Thesis Format applied successfully.")

if __name__ == "__main__":
    apply_whut_format(r"d:\SwapU\docs\thesis\SwapU毕业论文_正式定稿版.docx", r"d:\SwapU\docs\thesis\SwapU毕业论文_最终标准格式版.docx")
